from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from skp_core import rupiah, table_for_export


def build_docx(data: dict[str, Any], calc: dict[str, Any]) -> bytes:
    doc = Document(); doc.styles["Normal"].font.size = Pt(9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT KONFIRMASI PREMI" + (" - ENDORSEMENT" if data["mode"] == "Endorsement" else "")); run.bold = True; run.font.size = Pt(14)
    doc.add_paragraph(f"Nomor: {data['number']}\nTanggal: {data['letter_date']:%d-%m-%Y}\nKepada: {data['to']}\nTertanggung: {data['insured']}\nAlamat: {data['insured_address']}")
    if data["mode"] == "Endorsement":
        doc.add_paragraph(f"Nomor Polis: {data['old_policy_no']}\nPeriode Polis Lama: {data['old_start']:%d-%m-%Y} s.d. {data['old_end']:%d-%m-%Y}\nEfektif Endorsement: {data['effective']:%d-%m-%Y}\nTanggal Akhir Setelah Endorsement: {data['new_end']:%d-%m-%Y}")
    else:
        doc.add_paragraph(f"Periode Pertanggungan: {data['start']:%d-%m-%Y} s.d. {data['end']:%d-%m-%Y}")
    doc.add_heading("Rincian Obyek", level=2)
    if data["mode"] == "Endorsement":
        headers = ["Obyek","Perubahan","SI Lama","SI Baru","Kode/Kelas Lama","Rate Lama","Kode/Kelas Baru","Rate Baru","Adjustment"]
        table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
        for i,h in enumerate(headers): table.rows[0].cells[i].text = h
        for r in calc["rows"]:
            vals=[r["Obyek"],r["Perubahan"],rupiah(r["SI Lama"]),rupiah(r["SI Baru"]),f"{r['Kode Lama']} / {r['Kelas Lama']}",f"{r['Rate Lama (‰)']:.3f}‰",f"{r['Kode Baru']} / {r['Kelas Baru']}",f"{r['Rate Baru (‰)']:.3f}‰",rupiah(r["Adjustment"])]
            cells=table.add_row().cells
            for i,v in enumerate(vals): cells[i].text=str(v)
    else:
        headers=["Obyek","SI","Okupasi","Kode","Kelas","Rate","Range OJK","Premi FLEXAS"]
        table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
        for i,h in enumerate(headers): table.rows[0].cells[i].text=h
        for r in calc["rows"]:
            vals=[r["Obyek"],rupiah(r["SI"]),r["Okupasi"],r["Kode"],r["Kelas"],f"{r['Rate (‰)']:.3f}‰",r["Range OJK"],rupiah(r["Premi FLEXAS"])]
            cells=table.add_row().cells
            for i,v in enumerate(vals): cells[i].text=str(v)
    doc.add_heading("Ringkasan Premi", level=2)
    t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
    for k,v in table_for_export(data,calc):
        cells=t.add_row().cells; cells[0].text=k; cells[1].text=v
    doc.add_paragraph("Catatan: rate okupasi merupakan rate dasar Property/FLEXAS. Premi perluasan yang memerlukan zona/tabel tersendiri tidak dihitung otomatis dalam master okupasi. Untuk endorsement, perubahan premi dihitung sejak tanggal efektif endorsement untuk sisa periode pertanggungan.")
    sign_name = str(data.get('signer', '')).strip() or '________________________'
    doc.add_paragraph(f"\nHormat kami,\nPT Asuransi Jasa Indonesia\n\n\n{sign_name}\nRO Manager")
    buf=BytesIO(); doc.save(buf); return buf.getvalue()


def build_pdf(data: dict[str, Any], calc: dict[str, Any]) -> bytes:
    buf=BytesIO(); styles=getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle",parent=styles["Title"],alignment=TA_CENTER,fontSize=13,leading=16))
    small=ParagraphStyle(name="Small",parent=styles["BodyText"],fontSize=6.6,leading=8)
    body=ParagraphStyle(name="Body9",parent=styles["BodyText"],fontSize=8.5,leading=11)
    story=[Paragraph("SURAT KONFIRMASI PREMI"+(" - ENDORSEMENT" if data["mode"]=="Endorsement" else ""),styles["CenterTitle"]),Spacer(1,.25*cm)]
    info=[["Nomor",data["number"]],["Tanggal",f"{data['letter_date']:%d-%m-%Y}"],["Kepada",data["to"]],["Tertanggung",data["insured"]],["Alamat",data["insured_address"]]]
    if data["mode"]=="Endorsement":
        info += [["Nomor Polis",data["old_policy_no"]],["Periode Lama",f"{data['old_start']:%d-%m-%Y} s.d. {data['old_end']:%d-%m-%Y}"],["Efektif Endorsement",f"{data['effective']:%d-%m-%Y}"],["Akhir Setelah Endorsement",f"{data['new_end']:%d-%m-%Y}"]]
    else: info += [["Periode",f"{data['start']:%d-%m-%Y} s.d. {data['end']:%d-%m-%Y}"]]
    ti=Table([[Paragraph(str(a),body),Paragraph(str(b),body)] for a,b in info],colWidths=[4*cm,13*cm]); ti.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.25,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("BACKGROUND",(0,0),(0,-1),colors.whitesmoke)]))
    story += [ti,Spacer(1,.3*cm),Paragraph("<b>Rincian Obyek</b>",body)]
    if data["mode"]=="Endorsement":
        headers=["Obyek","SI Lama","SI Baru","Kode/Kelas Lama","Rate","Kode/Kelas Baru","Rate","Adjustment"]
        rows=[headers]+[[r["Obyek"],rupiah(r["SI Lama"]),rupiah(r["SI Baru"]),f"{r['Kode Lama']} / {r['Kelas Lama']}",f"{r['Rate Lama (‰)']:.3f}‰",f"{r['Kode Baru']} / {r['Kelas Baru']}",f"{r['Rate Baru (‰)']:.3f}‰",rupiah(r["Adjustment"])] for r in calc["rows"]]
        widths=[2.4*cm,2.3*cm,2.3*cm,2.1*cm,1.2*cm,2.1*cm,1.2*cm,2.4*cm]
    else:
        headers=["Obyek","SI","Okupasi","Kode","Kelas","Rate","Range OJK","Premi"]
        rows=[headers]+[[r["Obyek"],rupiah(r["SI"]),r["Okupasi"],r["Kode"],r["Kelas"],f"{r['Rate (‰)']:.3f}‰",r["Range OJK"],rupiah(r["Premi FLEXAS"])] for r in calc["rows"]]
        widths=[2.2*cm,2.2*cm,3.0*cm,1.2*cm,1.4*cm,1.6*cm,2.4*cm,2.6*cm]
    table=Table([[Paragraph(str(v),small) for v in row] for row in rows],colWidths=widths,repeatRows=1); table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.25,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.whitesmoke),("VALIGN",(0,0),(-1,-1),"TOP")]))
    story += [table,Spacer(1,.3*cm),Paragraph("<b>Ringkasan Premi</b>",body)]
    summary=Table([[Paragraph(k,body),Paragraph(v,body)] for k,v in table_for_export(data,calc)],colWidths=[10.2*cm,6.8*cm]); summary.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.25,colors.grey),("ALIGN",(1,0),(1,-1),"RIGHT")]))
    story += [summary,Spacer(1,.3*cm),Paragraph("Rate okupasi adalah rate dasar Property/FLEXAS. Premi perluasan yang memerlukan zona/tabel tersendiri harus dihitung terpisah. Untuk endorsement, perubahan premi dihitung sejak tanggal efektif endorsement untuk sisa periode.",body),Spacer(1,.55*cm),Paragraph(f"Hormat kami,<br/>PT Asuransi Jasa Indonesia<br/><br/><br/>{(str(data.get('signer', '')).strip() or '________________________')}<br/>RO Manager",body)]
    doc=SimpleDocTemplate(buf,pagesize=A4,leftMargin=1.2*cm,rightMargin=1.2*cm,topMargin=1.0*cm,bottomMargin=1.0*cm); doc.build(story); return buf.getvalue()
